#!/usr/bin/env python3
"""Convert the hunting analysis markdown to a .docx file."""
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Read markdown
with open('/opt/data/falai_hunting_executivo_novos_negocios.md', 'r') as f:
    content = f.read()

lines = content.split('\n')

i = 0
while i < len(lines):
    line = lines[i]
    
    # Headers
    if line.startswith('# ') and not line.startswith('## '):
        p = doc.add_heading(line[2:], level=0)
    elif line.startswith('### '):
        p = doc.add_heading(line[4:], level=2)
    elif line.startswith('## '):
        p = doc.add_heading(line[3:], level=1)
    elif line.startswith('> '):
        # Blockquote
        p = doc.add_paragraph()
        run = p.add_run(line[2:])
        run.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)
    elif line.startswith('- ') or line.startswith('* '):
        # Bullet
        p = doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('|'):
        # Table
        rows = [line]
        j = i + 1
        while j < len(lines) and lines[j].startswith('|'):
            rows.append(lines[j])
            j += 1
        
        # Parse table
        table_data = []
        for row in rows:
            cells = [c.strip() for c in row.split('|')[1:-1]]
            table_data.append(cells)
        
        # Skip separator rows
        filtered = [r for r in table_data if not all(re.match(r'^[-: ]+$', c) for c in r)]
        
        if len(filtered) > 0:
            table = doc.add_table(rows=len(filtered), cols=len(filtered[0]))
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            for ri, row_data in enumerate(filtered):
                for ci, cell_text in enumerate(row_data):
                    cell = table.rows[ri].cells[ci]
                    # Clean markdown bold
                    cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                    cell.text = cell_text
                    if ri == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
        
        doc.add_paragraph()  # spacing after table
        i = j
        continue
    elif line.strip() == '---':
        # Horizontal rule
        p = doc.add_paragraph()
        run = p.add_run('─' * 40)
        run.font.color.rgb = RGBColor(200, 200, 200)
    elif line.strip():
        # Regular paragraph with bold handling
        p = doc.add_paragraph()
        # Handle **bold** and emoji
        parts = re.split(r'(\*\*.*?\*\*|✅|⚠️|❌|🔴|🟡|🟢|🥇|🥈|🥉|🎯|📊|🔍|📋|📎|🏁|⭐|🟠)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
    
    i += 1

output_path = '/opt/data/Hunting_Executivo_Novos_Negocios_Florianopolis.docx'
doc.save(output_path)
print(f'Document saved to {output_path}')